import easyocr
import os
import cv2 
from datetime import datetime
import pandas as pd 
import numpy as np
from difflib import get_close_matches
from email.message import EmailMessage
import ssl
import smtplib
import dill as loader
from email.mime.image import MIMEImage
from keys import mail_sender, sender_app_password
from geopy.geocoders import Nominatim
from docx import Document
from docx.shared import Inches
import random


ROOT_DIR = "./"

def closeMatches(patterns, word):
    matc = get_close_matches(word, patterns)[0]
    return matc
    
     
def make_doc(result_set, cat=0):

    print("Ticket Generation")
    
    df = pd.read_excel('database.xlsx', engine='openpyxl')
    rns = np.array(df.iloc[:, 0])
    
    for item in result_set:
        result = item
        try:
            matched = closeMatches(rns, result)
        except:
            continue
        
        match_row = df[df['Rno']==matched]
        
        now = datetime.now()

        doc = Document()

        doc.add_heading('Issued by Traffic Regulations Authority', 1)

        table = doc.add_table(rows=1, cols=2)
        
        data = [
            ['Registration Number', str(matched)],
            ['Name of Owner', match_row['Name']],
            ['Address', match_row['Address']],
            ['Date and Time', now.strftime("%d/%m/%Y %H:%M:%S")],
            ['Riding a motor cycle without helmet', '1000'],
            ['Total Fine Amount', '1000 INR']  
        ]
        
        if cat==1:
            data = [
                ['Registration Number', str(matched)],
                ['Name of Owner', match_row['Name']],
                ['Address', match_row['Address']],
                ['Date and Time', now.strftime("%d/%m/%Y %H:%M:%S")],
                ['Violation of Traffic Signal', '2000'],
                ['Total Fine Amount', '2000 INR']  
            ]

        if cat==2:
            data = [
                ['Registration Number', str(matched)],
                ['Name of Owner', match_row['Name']],
                ['Address', match_row['Address']],
                ['Date and Time', now.strftime("%d/%m/%Y %H:%M:%S")],
                ['Usage of CellPhone while Driving', '1500'],
                ['Total Fine Amount', '1500 INR']  
            ]
        
        row = table.rows[0].cells
        row[0].text = 'Specification'
        row[1].text = 'Value'
        
        for ref, val in data:
            row = table.add_row().cells
            row[0].text = ref 
            row[1].text = val
            
        if cat!=0: 
            doc.add_picture('./images/violation.jpg', width=Inches(6))

        fname = random.randint(1111, 9999)
        document_path = './challan/{0}.docx'.format(fname)
        doc.save(document_path)
        send_email_ticket_helmet(document_path, match_row['Email'])


def get_lat_long_from_address(address):
    locator = Nominatim(user_agent='myapp')
    location = locator.geocode(address)
    return location.latitude, location.longitude
    
def load_pkl_model():
    with open('./models/model.pkl', 'rb') as rf:
        mdl = loader.load(rf)

    return mdl

def perform_ocr(): 
    im_dir = ROOT_DIR+'images/' 
    reader = easyocr.Reader(['en'])
    fnames = os.listdir(im_dir)
    result_set = []
    for name in fnames:
        print(name)
        im = cv2.imread(im_dir+name)
        result = reader.readtext(im, paragraph="False")
        try:
            result = result[0][-1].replace(" ", '').upper()
            result_set.append([name, result])
            
        except:
            continue
  
    return result_set


def send_email_ticket_helmet(document_path, ticket_email):

    print("Email Generated")

    subject = 'Violation Ticket Generated'

    body = f'A ticket has been generated due to the caused traffic violation with the registered vehicle'

    em = EmailMessage()
    em['From'] = mail_sender
    em['To'] = ticket_email
    em['subject'] = subject
    em.set_content(body)

    with open(document_path, 'rb') as f:
        file_data = f.read()
        file_name = f.name
        em.add_attachment(file_data, maintype='application', subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_name)

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login(mail_sender, sender_app_password)
        smtp.send_message(em)




